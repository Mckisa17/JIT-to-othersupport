import pandas as pd
import docx
from docx.shared import Pt
from docx.shared import Inches

# Read the XLS file into a pandas dataframe
df = pd.read_excel('JIT.xls', skiprows=6, header=6)

# Put XLS in data frame
df = pd.DataFrame(df)
df = df[~df['PI (last, first)'].isin(['Active', 'Pending', 'Awarded'])]
df = df.reset_index(drop=True)

# for loop for determining sub award
for i in range(len(df)):
    if df.loc[i, 'Is Mayo Secondary?'] == 'N':
        df.loc[i, 'Is Mayo Secondary?'] = 'Mayo Clinic, Rochester, MN'
    else:
        df.loc[i, 'Is Mayo Secondary?'] = '(sub award of ____)'

# Data Clean
df['Period 1 Effort (Calendar months)'] = df['Period 1 Effort (Calendar months)'].round(2)
df['Period 2 Effort (Calendar months)'] = df['Period 2 Effort (Calendar months)'].round(2)
df['Period 3 Effort (Calendar months)'] = df['Period 3 Effort (Calendar months)'].round(2)
df['Period 4 Effort (Calendar months)'] = df['Period 4 Effort (Calendar months)'].round(2)
df['Period 5 Effort (Calendar months)'] = df['Period 5 Effort (Calendar months)'].round(2)
df['date_column'] = pd.to_datetime(df['Project Period End Date'], format='%m/%d/%Y')
df['Project Period Start Date'] = pd.to_datetime(df['Project Period Start Date'], format='%m/%d/%Y')
df['Project Period End Date'] = pd.to_datetime(df['Project Period End Date'], format='%m/%d/%Y')
df['Total Project Period Costs (Direct plus Indirect)'] = df['Total Project Period Costs (Direct plus Indirect)'].apply(lambda x: '${:,.2f}'.format(x))

# Create a new Word document
def create_word_document():
    doc = docx.Document()

    # Set font, size, and style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)
    doc.styles['Normal'].paragraph_format.line_spacing = 1.0
    doc.styles['Normal'].paragraph_format.space_after = Pt(0)

    # Loop over all items in dataframe
    for i in range(df.shape[0]):

        # Write content for each project
        doc.add_paragraph("*Title: " + df["Grant Title"].loc[i])
        doc.add_paragraph("Major Goals: NEED ")
        doc.add_paragraph("*Status of Support: "+ df["Status Category"].loc[i])
        doc.add_paragraph("Project Number: " + df["External Grant ID"].astype(str).loc[i])
        doc.add_paragraph("Name of PD/PI: " + df["PI (last, first)"].loc[i])
        doc.add_paragraph("*Source of Support: "+ df["Funding Agency"].loc[i])
        doc.add_paragraph("*Primary Place of Performance: "+ df["Is Mayo Secondary?"].loc[i])
        start_date = df['Project Period Start Date'].loc[i]
        end_date = df['Project Period End Date'].loc[i]
        doc.add_paragraph("Project/Proposal Start and End Date: (MM/YYYY) (if available): "+ start_date.strftime('%m/%d/%Y') + " - " + end_date.strftime('%m/%d/%Y'))
        doc.add_paragraph("* Total Award Amount (including Indirect Costs): "+ str(df["Total Project Period Costs (Direct plus Indirect)"].loc[i]))
        doc.add_paragraph("* Person Months (Calendar/Academic/Summer) per budget period.")
        doc.add_paragraph()

        for p in doc.paragraphs[-11:-1]:
            p.paragraph_format.space_after = Pt(6)

        # Add a table to the document and identify how many rows are needed
        rows = df['# of Budget Periods in Current Project Period'].astype(int).loc[i]
        table = doc.add_table(rows=rows +1, cols=2)
        table.style = 'Table Grid'
        doc.add_paragraph()

        # Add year values to the table based on the project end date for the specific grant
        df['year'] = df['date_column'].dt.year
        for row in range(rows+1):
            for col in range(2):
                if col == 0:
                    if row == 0:
                        table.cell(row, 0).text = 'Year (YYYY)'
                    else:
                        table.cell(row, 0).text = str(row)+'. '+str(int(df['year'][i] - rows + row))
                if col ==1:
                    if row == 0:
                        table.cell(row, 1).text = 'Person Months (##.##)'
                    else:
                        table.cell(row, 1).text = '{:.2f} calendar'.format(df['Period {} Effort (Calendar months)'.format(row)].loc[i])

        # Set the cell height and width
        for row in table.rows:
            for cell in row.cells:
                if col == 0:
                    cell.width = Inches(1.62)
                    cell.height = Inches(0.18)
                else:
                    cell.width = Inches(1.81)
                    cell.height = Inches(0.18)

        # Drop rows with NaN values
        for row in table.rows:
            if any('nan' in cell.text.lower() for cell in row.cells):
                table._element.remove(row._element)

    # Save the document
    doc.save('output1.docx')

create_word_document()
